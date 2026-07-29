from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/Custody_Tracker_TJ_Review.docx"


def set_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, bold=False, color="000000", before=0, after=8, line=1.15):
    set_spacing(paragraph, before, after, line)
    for run in paragraph.runs:
        set_run(run, size=size, bold=bold, color=color)


def add_para(doc, text="", size=11, bold=False, color="000000", before=0, after=8):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    set_spacing(p, before, after)
    return p


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run(run, size=26, bold=False)
    set_spacing(p, 0, 3)
    p = add_para(doc, subtitle, size=11, color="555555", before=0, after=18)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=20, bold=False)
    set_spacing(p, 20, 6)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=16, bold=False)
    set_spacing(p, 18, 6)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=14, bold=False, color="434343")
    set_spacing(p, 16, 4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run(run)
    set_spacing(p, 0, 4)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run(run)
    set_spacing(p, 0, 4)
    return p


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, bold=bold)
    set_spacing(p, 0, 4)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
    set_table_borders(table)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    add_para(doc, "", after=4)
    return table


def add_decision_box(doc):
    add_h2(doc, "Decisions Needed From TJ")
    decisions = [
        "Should 'My day -> Co-parent helped' ask about final overnight location before or after help details?",
        "Should 'Other parent's day -> I helped' use a help confirmation screen instead of the sleeping-location confirmation screen?",
        "For calls, pickups, and brief visits on the other parent's day, do we need to record which kid was involved?",
        "Should schedule-change context be captured only for overnight/custody deviations, or also for meaningful daytime changes?",
        "When should we begin backend planning for accounts, sync, screenshot storage, and privacy controls?"
    ]
    for item in decisions:
        add_bullet(doc, item)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    add_title(
        doc,
        "Custody Tracker Decision Tree Review",
        "Working document for TJ review: current flow, open product questions, reporting choices, and backend timing."
    )
    add_para(doc, "Prepared from the current local build and project notes. Last updated: June 9, 2026.", color="555555", after=14)

    add_decision_box(doc)

    add_h1(doc, "1. Current Product Position")
    add_para(doc, "The app is currently optimized around quick, credible daily custody logging. The strongest product constraint is that users should not be able to backfill unlimited history, because that could weaken trust in the log.")
    add_bullet(doc, "Users can backfill only yesterday.")
    add_bullet(doc, "Older empty calendar days show \"Nothing logged\" and remain read-only.")
    add_bullet(doc, "Schedule-change context is captured for true custody deviations, not ordinary calls, pickups, or brief visits.")
    add_bullet(doc, "Each saved entry records a loggedAt timestamp.")
    add_bullet(doc, "Entries can include an optional screenshot attachment for context.")

    add_h1(doc, "2. Current Check-In Flow")
    add_h2(doc, "Top-Level Choice")
    add_number(doc, "My day")
    add_number(doc, "Other parent's day")

    add_h2(doc, "My Day Branch")
    add_table(
        doc,
        ["Path", "What the app asks", "Where it ends"],
        [
            ["I had the kids", "Are all kids sleeping at your house? If no, select who is with you and where absent kids are.", "Kid location confirmation, diary, review, saved."],
            ["Co-parent helped", "Pick the kids co-parent helped with, then record what co-parent did for each kid.", "Returns to the sleeping-location question, then diary/review."],
            ["Kids ended up with co-parent", "Pick kids at co-parent's tonight, then answer schedule-change context questions.", "Diary, review, saved."]
        ],
        [2300, 4200, 2860]
    )

    add_h2(doc, "Other Parent's Day Branch")
    add_table(
        doc,
        ["Path", "What the app asks", "Where it ends"],
        [
            ["Co-parent had the kids", "Choose involvement: none, phone/FaceTime, drop-off or pickup, brief visit.", "Review, saved."],
            ["I helped", "Pick kids you helped with, then record what you did for each kid.", "Kid location confirmation, diary/review."],
            ["Kids ended up with me", "Pick kids who ended up with you, then answer schedule-change context questions.", "Diary, review, saved."]
        ],
        [2300, 4200, 2860]
    )

    add_h1(doc, "3. Schedule Change Context")
    add_para(doc, "This appears only when actual custody differs from the scheduled day.")
    add_bullet(doc, "My day -> kids ended up with co-parent.")
    add_bullet(doc, "Other parent's day -> kids ended up with me.")
    add_para(doc, "Questions captured:")
    add_number(doc, "Was this change agreed to in advance?")
    add_number(doc, "If yes, did you feel pressured to agree?")
    add_para(doc, "Stored fields: changeAgreed and changePressured. These display in Review, Calendar details, log/report views, and exported report text.")

    add_h1(doc, "4. Odd Paths To Review")
    add_table(
        doc,
        ["Flow", "Why it may feel odd", "Question for TJ"],
        [
            ["My day -> Co-parent helped", "The user records co-parent activity first, then returns to the sleeping-location question. It can feel like a loop.", "Should overnight location come before help details?"],
            ["Other parent's day -> I helped", "The flow ends at the same kid confirmation screen used for sleeping-location confirmation.", "Should this branch get a different confirmation screen focused on help/involvement?"],
            ["Other parent's day -> Co-parent had kids -> brief visit/call/pickup", "The app goes directly to Review without asking which kid was involved.", "Is lightweight logging enough, or should we record kid-level detail?"],
            ["Progress indicators", "Some branches have different step totals, and schedule-change branches feel longer.", "Should all branches use a simpler progress model?"]
        ],
        [2300, 3600, 3460]
    )

    add_h1(doc, "5. Calendar And Reporting Behavior")
    add_bullet(doc, "Calendar colors use the rule: color equals whose scheduled day it was; stripes mean something changed.")
    add_bullet(doc, "Soft blue: your scheduled day, kids with you.")
    add_bullet(doc, "Blue with diagonal stripes: your day, kids ended up with co-parent.")
    add_bullet(doc, "Soft green: other parent's scheduled day.")
    add_bullet(doc, "Green with diagonal stripes: other parent's day, kids ended up with you.")
    add_bullet(doc, "Calendar detail shows logged timestamp, change-context badges, screenshot thumbnail/full image, and entry details.")
    add_bullet(doc, "Reports currently include summary stats, date-range filters, custody deviation report first, pressure/agreed flags, print, and preview auto-scroll.")

    add_h1(doc, "6. Backend Timing")
    add_para(doc, "A backend should be evaluated soon, but not integrated deeply until the decision tree and report model are stable.")
    add_h2(doc, "Recommended Timing")
    add_number(doc, "Finish the local product shape: decision tree, report semantics, screenshot behavior, and calendar rules.")
    add_number(doc, "Write a clean data model for users, children, entries, child-level entry details, attachments, and exports.")
    add_number(doc, "Prototype Supabase first, with Firebase as a credible fallback.")
    add_number(doc, "Add auth and cloud sync behind a feature flag.")
    add_number(doc, "Only then migrate localStorage data.")

    add_h2(doc, "Likely Backend Shortlist")
    add_table(
        doc,
        ["Option", "Why it fits", "Watch-out"],
        [
            ["Supabase", "Postgres, auth, private storage, row-level security, edge functions. Strong fit for reporting and legal-style exports.", "Needs careful RLS/privacy design from day one."],
            ["Firebase", "Very mature auth, Firestore, storage, functions, hosting, app monitoring, and offline-friendly patterns.", "Report queries can become awkward in a document database."],
            ["Convex", "TypeScript-first backend with realtime sync and fast developer iteration.", "Need to validate report/export queries before committing."],
            ["Appwrite", "Open-source backend with auth, database, storage, and functions.", "Likely a second-choice unless self-hosting/open-source control becomes central."]
        ],
        [1800, 4600, 2960]
    )

    add_h1(doc, "7. Testing And Release Hygiene")
    add_para(doc, "Current tests added to prevent repeat regressions:")
    add_bullet(doc, "Contract test for missing inline handlers, required screens, CSS hooks, and versioned assets.")
    add_bullet(doc, "Browser smoke test for Settings nav, check-in card styling, split-night defaults, Brief Visit review state, schedule-change context, Reports preview/filter behavior, Calendar trends, custom labels, and long-name overflow.")
    add_para(doc, "Before each push:")
    add_number(doc, "Run node scripts/contract-test.mjs.")
    add_number(doc, "Run npm run test:smoke, or the bundled Playwright smoke command if npm dependencies are not installed.")
    add_number(doc, "Update TODO and release notes with product decisions and regressions fixed.")

    add_h1(doc, "8. Suggested Review Agenda")
    add_number(doc, "Confirm whether the decision tree matches how TJ wants to think about custody documentation.")
    add_number(doc, "Resolve the three odd paths listed above.")
    add_number(doc, "Decide what level of kid-specific detail is needed for brief visits, calls, and pickups.")
    add_number(doc, "Confirm backend timing and privacy expectations.")
    add_number(doc, "Choose the next product area to spec: Trends, Reports, or Settings.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
